from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
import holidays

def generate_date_range(start_date, weeks):
    date_format = "%m/%d"
    start_date = datetime.strptime(start_date, date_format)

    # Create a new Word document
    doc = Document()

    # Add a table to the document
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.autofit = False

    # Initialize US holidays
    us_holidays = holidays.UnitedStates(years=start_date.year)

    for week in range(weeks):
        end_date = start_date + timedelta(days=6)
        week_dates = [start_date + timedelta(days=i) for i in range(7)]

        # Check if any day within the week is a US holiday
        week_holidays = [us_holidays[date] for date in week_dates if date in us_holidays]
        
        date_range = f"{start_date.strftime(date_format)} - {end_date.strftime(date_format)}"

        # Add the holiday with date if any day within the week is a holiday
        if week_holidays:
            holiday_date = [date for date in week_dates if date in us_holidays][0]
            holiday_text = f"{holiday_date.strftime(date_format)}: {week_holidays[0]} (US Holiday)"
            date_range += f"\n{holiday_text}"

        # Insert a new row in the table
        row = table.add_row().cells
        cell_paragraph = row[0].paragraphs[0]
        
        # Set font for all text in the cell to Times New Roman size 11
        for run in cell_paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

            # # Set font color to red for holiday text
            # if "US Holiday" in run.text:
            #     run.font.color.theme_color = 1  # 1 corresponds to the 'dark red' theme color

        cell_paragraph.add_run(str(week + 1))
        cell_paragraph.add_run("\n" + date_range)

        start_date = end_date + timedelta(days=1)

    # Save the document
    doc.save('output.docx')

# User input
start_date_input = input("Enter the starting date (MM/DD): ")
weeks_to_generate = 10

# Generate and save the Word document
generate_date_range(start_date_input, weeks_to_generate)
print("saved as output.docx")
